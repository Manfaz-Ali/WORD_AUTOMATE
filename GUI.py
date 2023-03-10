import sys

from PyQt6.QtWidgets import *
from PyQt6.uic import loadUiType
import docx
from docx.shared import Inches,Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd
from docx.enum.style import WD_STYLE_TYPE
ui, _ = loadUiType("MAIN1.ui")





class MainApp(QMainWindow, ui):
    def __init__(self, parent=None):
        super(MainApp, self).__init__(parent)
        QMainWindow.__init__(self)
        self.doc_name = None
        self.s_grade = None
        self.setupUi(self)
        self.doc = docx.Document()
        self.set_pageMargin()
        self.button_actions()

    
    def button_actions(self):
        self.radioButton1.clicked.connect(self.radio_button_clicked)
        self.radioButton2.clicked.connect(self.radio_button_clicked)
        
        
        #-----------------------------------------------------------
        self.pushButton_Save.clicked.connect(self.save_doc)
        self.pushButton_HdrFtr.clicked.connect(self.set_HdrFtr)
        self.pushButton_PGraph.clicked.connect(self.para)
        #--------------------------------------------------------
        self.pushButton_upload_csv.clicked.connect(self.upload_csv)
        self.pushButton_Table_generate.clicked.connect(self.tabular)
        self.pushButton_upload_img.clicked.connect(self.upload_img)
        self.pushButton_Image_generate.clicked.connect(self.img)
        self.pushButton_TableHeading.clicked.connect(self.HeadingTable)
        self.pushButton.clicked.connect(self.save_LM)
        

        
        
        

    def save_doc(self):
        doc = self.doc
        doc.save("example.docx")

    def get_docRef(self):
        self.ref = self.lineEdit_DocRef.text()
        return self.ref


    def get_LmRef(self):
        ref = self.lineEdit_LM_NO.text()
        
        return ref

    def get_docRev(self):
        self.rev = self.lineEdit_DocRev.text()
        return self.rev
    
    def get_docDate(self):
        self.date = self.lineEdit_DocDate.text()
        return self.date

    def get_LmDate(self):
        date = self.lineEdit_Date.text()
        return date

    def get_To(self):
        to = self.lineEdit_To.text()
        return to
    
    
    def add_table(self,df):
        doc = self.doc
        table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
        table.style = "Table Grid"
        for row in table.rows:
             for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for row in table.rows:
            row.height = Inches(0.5)
        header = table.rows[0].cells
        for i in range(df.shape[1]):
            header[i].text = df.columns[i]
            if i == 0:
                header[i].paragraphs[0].runs[0].font.bold = True
            header[i].paragraphs[0].runs[0].font.size = Pt(12)
            header[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            

            header[i].paragraphs[0].runs[0].font.name = 'Arial'
        for i in range(df.shape[0]):
            row = table.rows[i+1].cells
            for j in range(df.shape[1]):
                cell = row[j].add_paragraph(str(df.values[i,j]))
                cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.runs[0].font.size = Pt(12)
                cell.runs[0].font.name = 'Arial'
                
                if j == 0:
                    cell.runs[0].font.bold = True

    
    def upload_csv(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Open CSV File", ".", "CSV Files (*.csv)")
        if file_path:
            self.lineEdit_csv_FilePath.setText(file_path)


    def tbl_heading(self):
        t_head = self.lineEdit_table_heading.text()
        t_head = t_head.upper()
        return t_head

    def tabular(self):
        t=self.tbl_heading()
        self.i_heading_handler(t)
        self.draw_table()


    def draw_table(self):
        # Read CSV file using pandas
        file_path = self.lineEdit_csv_FilePath.text()
        if file_path:
            df = pd.read_csv(file_path)
            self.add_table(df)

    

    def radio_button_clicked(self):
        
        if self.radioButton1.isChecked():
            return 1
        if self.radioButton2.isChecked():
            return 2
        

    
    
    
    

    


    def get_document_name(self):
        self.doc_name = self.lineEdit_DocName.text()
        return self.doc_name

    def get_security_grade(self):
        self.s_grade = self.lineEdit_SGrade.text()
        return self.s_grade

    def get_LmSGRD(self):
        s_grade = self.lineEdit_Security_Grades.text()
        s_grade = s_grade.upper()
        return s_grade

    def get_LmForName(self):
        name = self.lineEdit_Name.text()
        return name

    def get_LmForRank(self):
        rank = self.lineEdit_Rank.text()
        return rank


    def LmHeader(self,headerText):
        doc = self.doc
        header = doc.sections[0].header
        paragraph = header.paragraphs[0]
        paragraph.add_run(headerText)
        # set the paragraph properties
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        paragraph.style.font.name = 'Arial'
        paragraph.style.font.size = docx.shared.Pt(12)
        paragraph.style.font.bold = False

    def LmFooter(self,footerText):
        doc = self.doc
        footer = doc.sections[0].footer
        paragraph = footer.paragraphs[0]
        paragraph.add_run(footerText)
        # set the paragraph properties
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        paragraph.style.font.name = 'Arial'
        paragraph.style.font.size = docx.shared.Pt(12)
        paragraph.style.font.bold = False

    def LmHdrFtr(self,grade):
        self.LmHeader(grade)
        self.LmFooter(grade)

    def get_GrpOfLmFor(self):
        grp = self.lineEdit_Group.text()
        return grp

    def get_tellOfLmFor(self):
        tel = self.lineEdit_Tell.text()
        return tel

    def LmMaker(self,To,LM_REFRENCE,DATE,NAME,RANK,GROUP,TEL,SUB,para):
        To = To.upper()
        LM_REFRENCE = LM_REFRENCE.upper()
        DATE = DATE.upper()
        NAME = NAME.upper()
        RANK = RANK.upper()
        GROUP = GROUP.upper()
        TEL = TEL.upper()


        doc = self.doc
        doc.add_paragraph("")
        doc.add_paragraph("")
        #your_para = self.user_func()
        # Add the title
        org_name = "Avionics Production Factory".upper()
        title = doc.add_paragraph(org_name)
        title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        title.style.font.bold = True
        title.style.font.name = "Arial"
        title.style.font.size = docx.shared.Pt(12)
        run = title.runs[0]
        run.font.bold = True

        sub_title = "(DDD)"
        stitle = doc.add_paragraph(sub_title)
        stitle.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        stitle.style.font.bold = False
        stitle.style.font.name = "Arial"
        stitle.style.font.size = docx.shared.Pt(12)
        

        # Add two empty lines
        doc.add_paragraph("")
        doc.add_paragraph("")

        # Add "Jnab"
        ref_to = doc.add_paragraph(To)
        ref_to.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

        # Add two empty lines
        doc.add_paragraph("")
        doc.add_paragraph("")

        # Add "This is start"
        if SUB == '':
            sub = "Subject".upper()
        else:
            sub = SUB.upper()
        start = doc.add_paragraph(sub)
        start.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        start.style.font.name = "Arial"
        run = start.runs[0]
        run.font.bold = True
        run.font.size = Pt(14)

         # Add two empty lines
        doc.add_paragraph("")
        doc.add_paragraph("")

        # Add dummy paragraph
        no = 1
        para = f'{no}.\t'+ para
        dummy = doc.add_paragraph(para)
        dummy.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
        start.style.font.bold = False
        start.style.font.name = "Arial"
        start.style.font.size = docx.shared.Pt(12)
        no+=1
        # Add five empty lines
        for i in range(5):
            doc.add_paragraph("")

        n_length = len(NAME)
        if n_length >= 4 and n_length <=8:
            pass
        name = doc.add_paragraph(f"\t\t\t\t\t( {NAME} )")
        name.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        run = name.runs[0]
        run.font.bold = True
        run.font.size = Pt(12)
        name.style.font.name = "Arial"

        rank = doc.add_paragraph(f"\t\t\t\t\t {RANK}")
        rank.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        rank.style.font.bold = False
        rank.style.font.name = "Arial"

        group = doc.add_paragraph(f"\t\t\t\t\t {GROUP}")
        group.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        group.style.font.bold = False
        group.style.font.name = "Arial"

        tell = doc.add_paragraph(f"\t\t\t\t\t {TEL} ")
        tell.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        tell.style.font.bold = False
        tell.style.font.name = "Arial"

        doc.add_paragraph("")

        lm_ref = doc.add_paragraph()
        lm_ref.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
        lm_ref.style.font.bold = False
        lm_ref.style.font.name = "Arial"

        lm_ref.add_run("LM No ")
        lm_ref.add_run(LM_REFRENCE)
        lm_ref.add_run(" ")
        lm_ref.add_run("dated ")
        lm_ref.add_run(DATE)

        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.line_spacing = 1.0

        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.space_after = 0
        
        
    def get_LmSubject(self):
        sb = self.lineEdit_lmSubject.text()  
        return sb

    def inser_paragraph(self):
        para = self.plainTextEdit_LmParagraph.toPlainText().strip()
        return para

    def save_LM(self):
        sg = self.get_LmSGRD()
        
        self.LmHdrFtr(sg)
        to = self.get_To()
        lm_no = self.get_LmRef()
        date = self.get_LmDate()
        nam = self.get_LmForName()
        rnk = self.get_LmForRank()
        grp = self.get_GrpOfLmFor()
        tel = self.get_tellOfLmFor()
        sub = self.get_LmSubject()
        para = self.inser_paragraph()
        self.LmMaker(to,lm_no,date,nam,rnk,grp,tel,sub,para)


    def select_to(self):
        ab = self.comboBox_to.current



    def set_HdrFtr(self):
        doc_name = self.get_document_name()
        doc_sgrade = self.get_security_grade()
        doc_ref_no = self.get_docRef()
        doc_rev_no = self.get_docRev()
        doc_date = self.get_docDate()

        self.set_header(doc_name,doc_sgrade)
        self.set_footer(doc_sgrade,doc_ref_no,doc_rev_no,doc_date)


    def set_pageMargin(self):
        doc = self.doc
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(1.5)
            section.right_margin = Inches(0.5)


    def get_table_sHd(self):
        s_hd = self.lineEdit_TableHeadS.text()
        s_hd = s_hd.upper()
        return s_hd


    def get_table_lHd(self):
        l_hd = self.lineEdit_TableHeadL.text()
        l_hd = l_hd.upper()
        return l_hd
    
    def HeadingTable(self):
        s = self.get_table_sHd()
        l = self.get_table_lHd()
        print(s)
        print(l)
        self.table_heading(s,l)



    def table_heading(self,s_hd,l_hd):
        doc = self.doc
        doc.add_paragraph("")

        table = doc.add_table(rows=2, cols=1)
        table.style = 'Table Grid'
        

        table.cell(0, 0).text = s_hd
        table.cell(1, 0).text = l_hd

        run = table.cell(0, 0).paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(12)

        run = table.cell(1, 0).paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(14)

        table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(1, 0).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.rows[0].height = Inches(0.5)
        table.rows[1].height = Inches(0.5)
        
        
        



    def set_header(self,your_docNm,your_sGrd):
        doc = self.doc
        header = doc.sections[0].header
        if your_docNm == '':
            ui_hdr_doc_nam = "my new document for test"
        else:
            ui_hdr_doc_nam = your_docNm
        if your_sGrd == '':
            ui_hdr_s_grd = "confidential"
        else:
            ui_hdr_s_grd = your_sGrd
        ui_hdr_doc_nam = ui_hdr_doc_nam.upper()
        ui_hdr_s_grd = ui_hdr_s_grd.upper()
        print(ui_hdr_doc_nam)
        print(ui_hdr_s_grd)
        table = header.add_table(rows=1, cols=3, width=Inches(8.01))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Resize the table to fit the header
        # for all rows
        for row in table.rows:
            for cell in row.cells:
                cell.width = docx.shared.Inches(2.67)
        # first cell
        first_cell = table.cell(0, 0)
        first_cell.text = ui_hdr_doc_nam
        self.set_cell_font(first_cell,12)
        
        
        
        # second cell
        second_cell = table.cell(0, 1)
        second_cell.text = ui_hdr_s_grd
        self.set_cell_font(second_cell,12)
        # third cell
        third_cell = table.cell(0, 2)
        run = third_cell.paragraphs[0].add_run()
        picture = run.add_picture("pic.jpg")
        picture.width = docx.shared.Inches(0.5)
        picture.height = docx.shared.Inches(0.25)
        third_cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        third_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # header from top
        doc.sections[0].header_distance = Inches(0.5)
        #header row height
        for row in table.rows:
            row.height = Inches(0.5)
        
        
        
        
        print("header done")
    
    def set_footer(self,sGrd,dRef,dRev,dDate):
        doc = self.doc
        # Set the footer of the document
        section = doc.sections[-1]
        footer = section.footer
        # Create a table with one row and three columns
        table = footer.add_table(rows=2, cols=4, width=Inches(8))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        # Add content to the cells
        cell1 = table.cell(0, 0)
        cell1.text = "Document Ref".upper()
        cell1.width = Inches(2.5)
        self.set_cell_font(cell1,12)

        cell2 = table.cell(0, 1)
        cell2.text = "Rev No".upper()
        cell2.width = Inches(1.5)
        self.set_cell_font(cell2,12)

        cell3 = table.cell(0, 2)
        cell3.text = "DATE"
        cell3.width = Inches(2)
        self.set_cell_font(cell3, 12)

        cell4 = table.cell(0, 3)
        cell4.text = 'Page'
        cell4.width = Inches(2.5)
        self.set_cell_font(cell4, 12)

        cell5 = table.cell(1, 0)
        if dRef == '':
            doc_Ref = 'Enter Doc Ref'
        else:
            doc_Ref = dRef
        doc_Ref = doc_Ref.upper()
        print(f"Ref No In Footer : {doc_Ref}")
        cell5.text = doc_Ref
        cell5.width = Inches(2)
        self.set_cell_font(cell5, 12)

        cell6 = table.cell(1, 1)
        if dRev == '':
            doc_Rev = 'Enter Rev No'
        else:
            doc_Rev = dRev
        doc_Rev = doc_Rev.upper()
        print(f"Rev No In Footer : {doc_Rev}")
        cell6.text = doc_Rev
        cell6.width = Inches(1.5)
        self.set_cell_font(cell6, 12)

        cell7 = table.cell(1, 2)
        if dDate == '':
            doc_Date = 'Enter Date'
        else:
            doc_Date = dDate
        doc_Date = doc_Date.upper()
        print(f"Date In Footer : {doc_Date}")
        cell7.text = doc_Date
        cell7.width = Inches(2)
        self.set_cell_font(cell7, 12)

        cell8 = table.cell(1, 3)
        run = cell8.paragraphs[0].add_run()
        field = OxmlElement('w:fldSimple')
        field.set(qn('w:instr'), 'PAGE')
        run._r.append(field)
        run.add_text(' of ')
        field = OxmlElement('w:fldSimple')
        field.set(qn('w:instr'), 'NUMPAGES')
        run._r.append(field)
        self.set_cell_font(cell8,12)
        paragraph = footer.add_paragraph()
        if sGrd == '':
            p = "confidential"
        else:
            p = sGrd
        p = p.upper()
        print(f"Security Grade In Footer : {p}")
        footer_run = paragraph.add_run(p)
        footer_run.bold = False
        footer_run.font.size = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Inches(0.2)
        paragraph_format.space_after = Inches(0.2)
        # header from top
        doc.sections[0].footer_distance = Inches(0.5)
        # footer row height
        for row in table.rows:
            row.height = Inches(0.5)
        print("Footer done")



    def get_P_heading(self):
        heading = self.lineEdit_paragraph_heading.text()
        heading = heading.upper()
        return heading


    def get_Img_heading(self):
        heading = self.lineEdit_img_heading.text()
        heading = heading.upper()
        return heading


    def p_heading_handler(self,heading):
        p_head = heading 
        doc = self.doc
        doc.add_paragraph("")
        heading = doc.add_paragraph(p_head)
        font = heading.style.font
        font.name = 'Arial'
        font.size = Pt(12)
        for run in heading.runs:
            run.font.bold = True
        heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    def i_heading_handler(self,heading):
        p_head = heading 
        doc = self.doc
        doc.add_paragraph("")
        heading = doc.add_paragraph(p_head)
        font = heading.style.font
        font.name = 'Arial'
        font.size = Pt(12)
        for run in heading.runs:
            run.font.bold = False
        heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        

    def para_text(self):
        P_text = self.textEdit_paraGraph_text.toPlainText()
        return P_text

    
    def img(self):
        img_heading = self.get_Img_heading()
        self.i_heading_handler(img_heading)
        self.draw_img()



    def para(self):
        your_heading = self.get_P_heading()
        self.p_heading_handler(your_heading)
        self.add_paragraph1()

    def add_paragraph1(self):
        doc = self.doc
        check = self.radio_button_clicked()
        your_para = self.para_text()
        if your_para == '':
            paragraph = "This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. "
        else:
            paragraph = your_para
        if check == 1:
            doc.add_paragraph("")
            doc.add_paragraph("")
            paragraph = '\t' + paragraph 
            doc.add_paragraph(paragraph).style.font.bold = False
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(12)
            font.bold = False
            doc.add_paragraph("")
            
        elif check==2:
            doc.add_paragraph("")
            paragraph = '\t' + paragraph 
            paragraph = doc.add_paragraph(paragraph)
            paragraph.style = 'List Number'
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(12)
            font.bold = False
            doc.add_paragraph("")
            doc.add_paragraph("")
        
            
            

            
            

            

            

    


    def set_cell_font(self,cell,size):
        #cell.paragraphs[0].runs[0].font.size = Pt(size)
        cell.paragraphs[0].style.font.size = Pt(size)
        cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].style.font.bold = False
        cell.paragraphs[0].style.font.name = 'Arial'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def upload_img(self):
        self.file_path, _ = QFileDialog.getOpenFileName(self, 'Select Image File', '', 'JPEG Files (*.jpg);;JPG Files (*.jpeg);;PNG Files (*.png)')
        if self.file_path:
            self.lineEdit_img_FilePath.setText(self.file_path)

    def draw_img(self):
        # Read CSV file using pandas
        file_path = self.lineEdit_img_FilePath.text()
        if file_path:
            self.insert_img(file_path)


    def insert_img(self,path_img):
        doc = self.doc
        doc.add_paragraph()
        doc.add_picture(path_img, width=Inches(4))
        inline_shape = doc.inline_shapes[-1]
        width, height = inline_shape.width, inline_shape.height
        aspect_ratio = int(height / width)
        new_width = Inches(2)
        new_height = aspect_ratio * new_width
        inline_shape.width = new_width
        inline_shape.height = new_height
        paragraph = doc.paragraphs[-1]
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    










def main():
    app = QApplication(sys.argv)
    window = MainApp()
    window.show()
    app.exec()


if __name__ == "__main__":
    main()