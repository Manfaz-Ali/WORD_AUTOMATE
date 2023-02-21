import docx
from PyQt6.QtCore import Qt
from PyQt6.QtWidgets import QApplication, QWidget, QTextEdit, QPushButton, QVBoxLayout, QMessageBox

class MyWindow(QWidget):
    def __init__(self):
        super().__init__()

        # Create the user interface
        self.text_box = QTextEdit(self)
        self.submit_button = QPushButton("Add Paragraph", self)
        self.submit_button.clicked.connect(self.insert_paragraph)

        # Add the user interface components to the layout
        layout = QVBoxLayout()
        layout.addWidget(self.text_box)
        layout.addWidget(self.submit_button)
        self.setLayout(layout)

    def insert_paragraph(self):
        # Get the user input from the text box
        new_text = self.text_box.toPlainText().strip()

        # Load the document
        doc = docx.Document()

        # Add two line breaks before the new paragraph
        doc.add_paragraph("")
        doc.add_paragraph("")

        # Add the new paragraph
        doc.add_paragraph(new_text)

        # Add two line breaks after the new paragraph
        doc.add_paragraph("")
        doc.add_paragraph("")

        # Save the document
        doc.save('my_document.docx')

        # Show a message box to confirm that the new paragraph has been added
        QMessageBox.information(self, "Success", "New paragraph added successfully.")

if __name__ == '__main__':
    app = QApplication([])
    window = MyWindow()
    window.show()
    app.exec()
