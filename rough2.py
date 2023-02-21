import docx
from PyQt6.QtCore import Qt
from PyQt6.QtWidgets import QApplication, QWidget, QTextEdit, QPushButton, QVBoxLayout, QDialog, QDialogButtonBox, QMessageBox

class AddParagraphDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)

        # Create the user interface
        self.text_box = QTextEdit(self)
        self.button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel, Qt.Horizontal, self)
        self.button_box.accepted.connect(self.accept)
        self.button_box.rejected.connect(self.reject)

        # Add the user interface components to the layout
        layout = QVBoxLayout()
        layout.addWidget(self.text_box)
        layout.addWidget(self.button_box)
        self.setLayout(layout)

        self.setWindowTitle("Add Paragraph")

    def get_paragraph_text(self):
        return self.text_box.toPlainText().strip()

class MyWindow(QWidget):
    def __init__(self):
        super().__init__()

        # Create the user interface
        self.text_box = QTextEdit(self)
        self.submit_button = QPushButton("Add Paragraph", self)
        self.submit_button.clicked.connect(self.add_paragraph)

        # Add the user interface components to the layout
        layout = QVBoxLayout()
        layout.addWidget(self.text_box)
        layout.addWidget(self.submit_button)
        self.setLayout(layout)

    def add_paragraph(self):
        # Show the dialog to get the new paragraph from the user
        dialog = AddParagraphDialog(self)
        result = dialog.exec()

        if result == QDialog.Accepted:
            # Get the user input from the dialog
            new_text = dialog.get_paragraph_text()

            # Load the document
            doc = docx.Document('my_document.docx')

            # Add two line breaks before the new paragraph
            doc.add_paragraph("")
            doc.add_paragraph("")

            # Add the new paragraph
            doc.add_paragraph(new_text)

            # Add two line breaks after the new paragraph
            doc.add_paragraph("")
            doc.add_paragraph("")

            # Save the document
            doc.save('my_document.docx')

            # Show a message box to confirm that the new paragraph has been added
            QMessageBox.information(self, "Success", "New paragraph added successfully.")

            # Clear the user input from the text box in the dialog
            dialog.text_box.clear()

if __name__ == '__main__':
    app = QApplication([])
    window = MyWindow()
    window.show()
    app.exec()
