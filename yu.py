import docx

doc = docx.Document()

# Add two empty lines
doc.add_paragraph("")
doc.add_paragraph("")

# Add the title
title = doc.add_paragraph("Avionics Production Factory")
title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
title.style.font.bold = True
title.style.font.name = "Arial"
title.style.font.size = docx.shared.Pt(14)
title.style.font.all_caps = True

stitle = doc.add_paragraph("(DDD)")
stitle.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
stitle.style.font.bold = False
stitle.style.font.name = "Arial"
stitle.style.font.size = docx.shared.Pt(12)
stitle.style.font.all_caps = True

# Add two empty lines
doc.add_paragraph("")
doc.add_paragraph("")

# Add "Jnab"
jnab = doc.add_paragraph("Jnab")
jnab.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

# Add two empty lines
doc.add_paragraph("")
doc.add_paragraph("")

# Add "This is start"
start = doc.add_paragraph("SUBJECT HEADING")
start.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
start.style.font.bold = True
start.style.font.name = "Arial"
start.style.font.size = docx.shared.Pt(14)

# Add two empty lines
doc.add_paragraph("")
doc.add_paragraph("")

# Add dummy paragraph
dummy = doc.add_paragraph("Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed nec velit feugiat, ultricies tellus vitae, tincidunt tellus. Nullam vel tellus eu massa porttitor ullamcorper. Praesent vel est eleifend, faucibus leo eget, tincidunt enim. Morbi tristique, justo et dapibus pulvinar, velit odio euismod est, ac vestibulum augue justo at enim.")
dummy.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
start.style.font.bold = False
start.style.font.name = "Arial"
start.style.font.size = docx.shared.Pt(12)
# Add five empty lines
for i in range(5):
    doc.add_paragraph("")

# Add "rtyu", "rtwr", and "tyuue"
rtyu = doc.add_paragraph("( MANFAZ")
rtyu.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
rtyu.style.font.bold = False
rtyu.style.font.name = "Arial"

rtwr = doc.add_paragraph("HSM )")
rtwr.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
rtwr.style.font.bold = False
rtwr.style.font.name = "Arial"

tyuue = doc.add_paragraph("LM No")
tyuue.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
tyuue.style.font.bold = False
tyuue.style.font.name = "Arial"

tyuue.add_run("")
tyuue.add_run("Date")



# Save the document
doc.save("my_doc.docx")
